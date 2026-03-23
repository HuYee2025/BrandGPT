"""
品牌AI顾问系统 - Flask主应用
"""
import os
import json
from flask import Flask, render_template, request, redirect, url_for, flash, session, jsonify, Response
from flask_login import LoginManager, login_user, logout_user, login_required, current_user
from datetime import datetime
from config import config
from models import db, User, Conversation, Message, Report, ModelConfig
from ai_engine import (
    SYSTEM_PROMPT,
    generate_question_prompt,
    generate_report_prompt,
    get_initial_question,
    get_model_client
)

app = Flask(__name__,
            template_folder=os.path.join(os.path.dirname(os.path.abspath(__file__)), 'templates'),
            static_folder=os.path.join(os.path.dirname(os.path.abspath(__file__)), 'static'))
app.config.from_object(config[os.getenv('FLASK_ENV', 'default')])

# 初始化扩展
db.init_app(app)
login_manager = LoginManager(app)
login_manager.login_view = 'login'
login_manager.login_message = '请先登录'


@login_manager.user_loader
def load_user(user_id):
    return User.query.get(int(user_id))


# 创建数据库
with app.app_context():
    db.create_all()


# ==================== 路由 ====================

@app.route('/')
def index():
    """首页"""
    if current_user.is_authenticated:
        return redirect(url_for('chat'))
    return render_template('index.html')


@app.route('/architecture')
def architecture():
    """系统架构演示页面"""
    return render_template('architecture.html')


@app.route('/tools')
def tools():
    """经营小工具页面"""
    return render_template('tools.html')


@app.route('/report-demo')
def report_demo():
    """报告示例页面"""
    return render_template('report_demo.html')


@app.route('/report-demo/download')
def download_report_demo():
    """下载示例报告Word文档"""
    from flask import send_file
    from docx import Document
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from io import BytesIO

    doc = Document()

    # 设置默认字体为微软雅黑
    style = doc.styles['Normal']
    style.font.name = 'Microsoft YaHei'
    style.font.size = Pt(11)
    # 设置中文字体的东亚字符集
    style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Microsoft YaHei')
    style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Microsoft YaHei')
    style._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Microsoft YaHei')

    # 设置所有内置样式
    for style_name in ['Heading 1', 'Heading 2', 'Heading 3', 'List Bullet', 'List Paragraph']:
        if style_name in doc.styles:
            s = doc.styles[style_name]
            s.font.name = 'Microsoft YaHei'
            s.font.size = Pt(11)
            s._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', 'Microsoft YaHei')
            s._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', 'Microsoft YaHei')
            s._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', 'Microsoft YaHei')

    # 辅助函数：设置run的字体
    def set_run_font(run, font_name='Microsoft YaHei', font_size=None):
        run.font.name = font_name
        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}ascii', font_name)
        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}hAnsi', font_name)
        run._element.rPr.rFonts.set('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}eastAsia', font_name)
        if font_size:
            run.font.size = font_size

    # 标题
    title = doc.add_heading('「面咖」面馆+咖啡商业计划书', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        set_run_font(run, 'Microsoft YaHei', Pt(22))

    # 内容
    content = """
一、项目概述

「面咖」是一家创新型餐饮品牌，将传统面食文化与现代咖啡生活方式相结合，打造"一碗好面+一杯好咖啡"的复合式餐饮空间。项目面向25-40岁的都市白领群体，提供快捷、优质、性价比高的中式面食与精品咖啡服务。

1.1 项目定位

品牌理念：传统与时尚的完美融合
核心产品：手工拉面 + 精品咖啡
目标客群：都市白领、年轻家庭、文艺青年
消费场景：早餐、午餐、下午茶、商务简餐

1.2 项目优势

差异化竞争：避免与传统面馆和咖啡馆的正面竞争
提高客单价：双品类带来更高的客均消费
时段互补：咖啡提升早餐和下午茶时段的营收
场景丰富：满足一人食、朋友小聚、商务洽谈等多种场景

二、市场分析

2.1 行业现状

根据美团研究院数据，2024年中国面食市场规模突破5000亿元，年增长率保持在8%以上。咖啡市场规模超过2000亿元，连续多年保持15%以上的增速。两者的消费人群高度重合，市场基础扎实。

2.2 目标市场

以一线城市核心商圈和二线城市繁华地段为主，目标客群特征：
年龄：25-40岁，占比70%以上
月收入：8000-20000元
消费习惯：注重品质与效率，愿意为体验付费

2.3 竞争分析

竞争对手分析：
传统面馆 - 优势：品类丰富、价格优势；劣势：环境简陋、缺乏品牌
连锁咖啡店 - 优势：品牌影响力、空间体验；劣势：主食选择有限、价格偏高
日式拉面店 - 优势：产品标准化、品牌精致；劣势：口味单一、翻台率低

三、产品与服务

3.1 产品结构

品类 | 主打产品 | 价格区间 | 占比
招牌面食 | 红烧牛肉面、担担面、酸汤面 | 28-38元 | 45%
精品咖啡 | 拿铁、美式、手冲 | 22-38元 | 25%
小食甜点 | 卤味、凉菜、芝士蛋糕 | 12-28元 | 20%
饮品配套 | 鲜榨果汁、气泡水 | 15-22元 | 10%

3.2 服务特色

快速出餐：面食3分钟、咖啡2分钟，确保高峰时段效率
半自助服务：线上点单、扫码支付，减少等待时间
会员体系：积分兑换、会员专享价、生日礼遇
社群运营：定期组织品鉴会、烹饪课堂等活动

四、选址与装修

4.1 选址标准

首选位置：写字楼底商、商业综合体、地铁交通枢纽
面积要求：80-150平方米
层高要求：3.5米以上，具备餐饮条件
租金预算：月租金控制在营收的10%以内

4.2 装修风格

采用"现代中式"风格，融合工业风与国潮元素：
开放式厨房，可视化制作过程
原木色桌椅搭配皮质卡座
墙面采用手绘插画，讲述面食文化故事
暖色调灯光营造温馨氛围

五、盈利预测

5.1 投资预算

项目 | 金额 | 备注
品牌加盟费 | 8万元 | 含品牌使用、技术培训
装修费用 | 15万元 | 100㎡中等装修
设备采购 | 12万元 | 厨房设备、咖啡设备
首月租金 | 3万元 | 押一付三
原材料储备 | 5万元 | 首批原材料采购
流动资金 | 7万元 | 日常运营周转
合计 | 50万元 |

5.2 营收预测

月份 | 日均客流 | 客单价 | 月营收
开业期(1-3月) | 80人 | 45元 | 10.8万元
成长期(4-6月) | 120人 | 50元 | 18万元
稳定期(7-12月) | 150人 | 55元 | 24.75万元

5.3 成本结构

成本项目 | 占比 | 说明
原材料成本 | 30% | 食材、调料、咖啡豆
人工成本 | 25% | 员工工资、社保
租金成本 | 10% | 月租3万元
水电杂费 | 5% | 能耗、日常损耗
营销费用 | 5% | 推广、优惠活动
其他支出 | 5% | 设备折旧、税费
净利润率 | 20% |

六、发展规划

6.1 三年发展目标

第一年：打磨单店模型，建立标准化体系，开出2-3家直营店
第二年：开放加盟，发展5-8家加盟店，覆盖核心城市
第三年：品牌规模化运营，开出20家以上门店，布局全国市场

6.2 品牌愿景

成为中国新式餐饮的领先品牌，让"面咖"成为都市人生活中不可或缺的品质之选。

七、风险与应对

7.1 主要风险

选址风险：位置不佳导致客流不足
运营风险：人员流动大，品质难以保证
竞争风险：同类型竞争者模仿跟进
供应链风险：原材料价格波动影响成本

7.2 应对策略

建立专业的选址评估体系，严格把关
完善员工培训体系，建立激励机制
持续创新，保持产品差异化优势
建立稳定的供应商合作关系
"""

    # 解析内容并添加到文档
    table = None
    lines = content.strip().split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if line.startswith('#'):
            level = min(line.count('#'), 3)
            h = doc.add_heading(line.replace('#', '').strip(), level=level)
            for run in h.runs:
                set_run_font(run)
        elif line.startswith('-') or line.startswith('•'):
            p = doc.add_paragraph(line.lstrip('-• ').strip(), style='List Bullet')
            for run in p.runs:
                set_run_font(run)
        elif ' | ' in line:
            # 表格行（以 | 分隔）
            parts = [p.strip() for p in line.split('|')]
            parts = [p for p in parts if p]  # 过滤空字符串
            if parts:
                if '占比' in parts or '项目' in parts or '月份' in parts or '竞争对手' in parts or '品类' in parts or '成本项目' in parts:
                    # 表头
                    table = doc.add_table(rows=1, cols=len(parts))
                    table.style = 'Table Grid'
                    hdr_cells = table.rows[0].cells
                    for i, part in enumerate(parts):
                        hdr_cells[i].text = part
                        for run in hdr_cells[i].paragraphs[0].runs:
                            set_run_font(run)
                            run.font.bold = True
                else:
                    if table is not None:
                        row_cells = table.add_row().cells
                        for i, part in enumerate(parts):
                            if i < len(row_cells):
                                row_cells[i].text = part
                                for run in row_cells[i].paragraphs[0].runs:
                                    set_run_font(run)
        else:
            p = doc.add_paragraph(line)
            for run in p.runs:
                set_run_font(run)

    # 添加免责声明
    doc.add_paragraph('')
    p = doc.add_paragraph('——' * 20)
    for run in p.runs:
        set_run_font(run)
    disclaimer = doc.add_paragraph('本文由 AI 模型生成，未必正确无误，请核查重要信息。')
    disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in disclaimer.runs:
        set_run_font(run)

    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name='「面咖」商业计划书.docx',
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@app.route('/复利计算器.html')
def fuli_calc():
    """复利计算器页面"""
    return render_template('复利计算器.html')


@app.route('/餐饮店利润率计算器.html')
def profit_calc():
    """餐饮店利润率计算器页面"""
    return render_template('餐饮店利润率计算器.html')


# ==================== 用户认证 ====================

@app.route('/register', methods=['GET', 'POST'])
def register():
    """用户注册"""
    if request.method == 'POST':
        username = request.form.get('username')
        email = request.form.get('email')
        password = request.form.get('password')

        # 检查用户是否存在
        if User.query.filter_by(username=username).first():
            flash('用户名已存在', 'error')
            return redirect(url_for('register'))

        if User.query.filter_by(email=email).first():
            flash('邮箱已被注册', 'error')
            return redirect(url_for('register'))

        # 创建用户
        user = User(username=username, email=email)
        user.set_password(password)

        # 设置默认API Key（如果有环境变量）
        if os.getenv('OPENAI_API_KEY'):
            user.api_key = os.getenv('OPENAI_API_KEY')
            user.model_provider = 'openai'
            user.model_name = 'gpt-4'

        db.session.add(user)
        db.session.commit()

        flash('注册成功，请登录', 'success')
        return redirect(url_for('login'))

    return render_template('register.html')


@app.route('/login', methods=['GET', 'POST'])
def login():
    """用户登录"""
    if request.method == 'POST':
        username = request.form.get('username')
        password = request.form.get('password')

        user = User.query.filter_by(username=username).first()

        if user and user.check_password(password):
            login_user(user)
            flash('登录成功', 'success')
            next_page = request.args.get('next')
            return redirect(next_page or url_for('chat'))
        else:
            flash('用户名或密码错误', 'error')

    return render_template('login.html')


@app.route('/logout')
@login_required
def logout():
    """退出登录"""
    logout_user()
    flash('已退出登录', 'info')
    return redirect(url_for('login'))


# ==================== AI问答 ====================

@app.route('/chat')
def chat():
    """AI问答页面"""
    prompt = request.args.get('prompt', '')
    return render_template('chat.html', initial_prompt=prompt)


@app.route('/chat', methods=['POST'])
def chat_post():
    """处理聊天消息"""
    message = request.form.get('message', '')
    is_initial = request.form.get('is_initial', 'false') == 'true'

    # 如果是初始请求（没有选择品类），返回开场白
    if is_initial:
        initial_prompt = """你好！很高兴认识你！

我是一名积累了丰富实践商业经验的创业顾问，会从专业的角度提出问题，逐渐将你头脑中可能有些杂乱的想法梳理清楚。

首先，请告诉我，您想在什么地方做什么样的生意呢？希望你能尽可能的表述详细，不要在意语句是否通顺。

让我们从杂乱中寻找到规律和机会，解开您的财富密码。"""
        return jsonify({'response': initial_prompt})

    if not message:
        return jsonify({'error': '消息不能为空'}), 400

    try:
        # 获取或创建临时会话
        if 'temp_context' not in session:
            session['temp_context'] = []

        # 检查是否是第一次对话
        is_first_message = len(session.get('temp_context', [])) == 0

        # 构建消息列表
        messages = [{'role': 'system', 'content': SYSTEM_PROMPT}] + session['temp_context'] + [{'role': 'user', 'content': message}]

        # 读取配置文件（从项目外部读取）
        config_path = os.path.join(os.path.dirname(__file__), '..', 'api_config.json')
        api_key = ''
        provider = 'openai'
        model_name = 'gpt-4'
        api_url = ''

        if os.path.exists(config_path):
            with open(config_path, 'r', encoding='utf-8') as f:
                config = json.load(f)
                api_key = config.get('api_key', '')
                provider = config.get('provider', 'openai')
                model_name = config.get('model_name', 'gpt-4')
                api_url = config.get('api_url', '')

        # 如果配置文件没有，使用环境变量
        if not api_key:
            api_key = os.getenv('OPENAI_API_KEY', '')
        if not model_name:
            model_name = os.getenv('AI_MODEL', 'gpt-4')
        if not api_url:
            api_url = os.getenv('AI_API_URL', '')

        if not api_key:
            return jsonify({'error': '请先在API管理中配置API Key'}), 400

        # 调用AI，传递base_url
        client = get_model_client(provider, api_key, model_name, base_url=api_url if api_url else None)
        response = client.chat(messages)

        # 更新上下文
        session['temp_context'] = session.get('temp_context', []) + [
            {'role': 'user', 'content': message},
            {'role': 'assistant', 'content': response}
        ]

        # 如果是第一次对话，生成标题
        result = {'response': response}
        if is_first_message or (session.get('temp_context', []) == [] and message.startswith('我想开一家')):
            # 尝试从用户消息中提取类别名称
            import re
            match = re.search(r'我想开一家(.+?)[，,]', message)
            if match:
                # 直接使用类别名作为标题
                title = match.group(1)[:12]
            else:
                # 如果提取不到，调用AI生成标题
                title_prompt = f"请根据用户的以下问题提炼出12个字以内的关键词标题，不要有标点符号，直接返回标题：{message}"
                title_messages = [{'role': 'user', 'content': title_prompt}]
                title_response = client.chat(title_messages)
                # 清理标题（去除标点、空白等）
                title = title_response.strip().replace('。', '').replace('，', '').replace('！', '').replace('？', '').replace(' ', '')[:12]
            result['title'] = title

        return jsonify(result)

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/chat/clear-context', methods=['POST'])
def clear_chat_context():
    """清除聊天上下文（当用户清空所有历史记录时调用）"""
    if 'temp_context' in session:
        session.pop('temp_context', None)
    return jsonify({'success': True})


@app.route('/api/chat/new', methods=['POST'])
@login_required
def new_conversation():
    """创建新对话"""
    conversation = Conversation(
        user_id=current_user.id,
        title='新对话'
    )
    db.session.add(conversation)
    db.session.commit()

    # 添加系统消息
    system_msg = Message(
        conversation_id=conversation.id,
        role='system',
        content=SYSTEM_PROMPT
    )
    db.session.add(system_msg)

    # 添加初始问题
    initial_question = get_initial_question()
    assistant_msg = Message(
        conversation_id=conversation.id,
        role='assistant',
        content=initial_question
    )
    db.session.add(assistant_msg)
    db.session.commit()

    return jsonify({
        'conversation_id': conversation.id,
        'message': initial_question
    })


@app.route('/api/chat/<int:conversation_id>/send', methods=['POST'])
@login_required
def send_message(conversation_id):
    """发送消息"""
    data = request.get_json()
    user_message = data.get('message', '')

    conversation = Conversation.query.get_or_404(conversation_id)
    if conversation.user_id != current_user.id:
        return jsonify({'error': '无权访问'}), 403

    # 保存用户消息
    user_msg = Message(
        conversation_id=conversation.id,
        role='user',
        content=user_message
    )
    db.session.add(user_msg)
    db.session.commit()

    # 获取对话历史
    messages = Message.query.filter_by(conversation_id=conversation.id)\
        .order_by(Message.created_at).all()

    # 构建消息列表
    chat_messages = []
    for msg in messages:
        chat_messages.append({
            'role': msg.role,
            'content': msg.content
        })

    # 获取模型客户端
    api_key = current_user.api_key or os.getenv('OPENAI_API_KEY', '')
    provider = current_user.model_provider or 'openai'
    model_name = current_user.model_name or 'gpt-4'

    if not api_key:
        return jsonify({'error': '请先配置API Key'}), 400

    try:
        client = get_model_client(provider, api_key, model_name)

        # 调用AI
        response = client.chat(chat_messages, stream=False)

        # 保存AI响应
        assistant_msg = Message(
            conversation_id=conversation.id,
            role='assistant',
            content=response
        )
        db.session.add(assistant_msg)

        # 更新对话时间
        conversation.updated_at = datetime.utcnow()
        db.session.commit()

        return jsonify({
            'message': response,
            'conversation_id': conversation.id
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


@app.route('/api/chat/<int:conversation_id>/stream', methods=['POST'])
@login_required
def stream_chat(conversation_id):
    """流式聊天"""
    data = request.get_json()
    user_message = data.get('message', '')

    conversation = Conversation.query.get_or_404(conversation_id)
    if conversation.user_id != current_user.id:
        return jsonify({'error': '无权访问'}), 403

    # 保存用户消息
    user_msg = Message(
        conversation_id=conversation.id,
        role='user',
        content=user_message
    )
    db.session.add(user_msg)
    db.session.commit()

    # 获取对话历史
    messages = Message.query.filter_by(conversation_id=conversation.id)\
        .order_by(Message.created_at).all()

    # 构建消息列表
    chat_messages = []
    for msg in messages:
        chat_messages.append({
            'role': msg.role,
            'content': msg.content
        })

    # 获取模型客户端
    api_key = current_user.api_key or os.getenv('OPENAI_API_KEY', '')
    provider = current_user.model_provider or 'openai'
    model_name = current_user.model_name or 'gpt-4'

    if not api_key:
        return jsonify({'error': '请先配置API Key'}), 400

    def generate():
        try:
            client = get_model_client(provider, api_key, model_name)
            full_response = ''

            for chunk in client.chat(chat_messages, stream=True):
                full_response += chunk
                yield f"data: {json.dumps({'chunk': chunk})}\n\n"

            # 保存完整响应
            assistant_msg = Message(
                conversation_id=conversation.id,
                role='assistant',
                content=full_response
            )
            db.session.add(assistant_msg)

            conversation.updated_at = datetime.utcnow()
            db.session.commit()

            yield f"data: {json.dumps({'done': True})}\n\n"

        except Exception as e:
            yield f"data: {json.dumps({'error': str(e)})}\n\n"

    return Response(generate(), mimetype='text/event-stream')


@app.route('/api/conversations')
@login_required
def get_conversations():
    """获取对话列表"""
    conversations = Conversation.query.filter_by(user_id=current_user.id)\
        .order_by(Conversation.updated_at.desc()).all()

    return jsonify([{
        'id': c.id,
        'title': c.title,
        'updated_at': c.updated_at.strftime('%Y-%m-%d %H:%M')
    } for c in conversations])


@app.route('/api/conversations/<int:conversation_id>')
@login_required
def get_conversation(conversation_id):
    """获取对话详情"""
    conversation = Conversation.query.get_or_404(conversation_id)
    if conversation.user_id != current_user.id:
        return jsonify({'error': '无权访问'}), 403

    messages = Message.query.filter_by(conversation_id=conversation.id)\
        .order_by(Message.created_at).all()

    return jsonify([{
        'id': m.id,
        'role': m.role,
        'content': m.content,
        'created_at': m.created_at.strftime('%Y-%m-%d %H:%M')
    } for m in messages])


# ==================== 报告生成 ====================

@app.route('/report/<int:report_id>')
@login_required
def view_report(report_id):
    """查看报告"""
    report = Report.query.get_or_404(report_id)
    if report.user_id != current_user.id:
        return redirect(url_for('chat'))

    return render_template('report.html', report=report)


@app.route('/api/report/generate/<int:conversation_id>', methods=['POST'])
@login_required
def generate_report(conversation_id):
    """生成报告"""
    conversation = Conversation.query.get_or_404(conversation_id)
    if conversation.user_id != current_user.id:
        return jsonify({'error': '无权访问'}), 403

    # 获取所有消息
    messages = Message.query.filter_by(conversation_id=conversation.id)\
        .order_by(Message.created_at).all()

    # 构建对话历史
    conversation_history = []
    for msg in messages:
        if msg.role in ['user', 'assistant']:
            conversation_history.append({
                'role': msg.role,
                'content': msg.content
            })

    # 获取模型客户端
    api_key = current_user.api_key or os.getenv('OPENAI_API_KEY', '')
    provider = current_user.model_provider or 'openai'
    model_name = current_user.model_name or 'gpt-4'

    if not api_key:
        return jsonify({'error': '请先配置API Key'}), 400

    try:
        client = get_model_client(provider, api_key, model_name)

        # 生成报告
        prompt = generate_report_prompt(conversation_history)
        report_content = client.chat([
            {'role': 'system', 'content': '你是一个专业的商业计划书撰写专家。请生成专业、详实的商业计划书。'},
            {'role': 'user', 'content': prompt}
        ])

        # 保存报告
        report = Report(
            user_id=current_user.id,
            conversation_id=conversation.id,
            title=f"商业计划书 - {datetime.now().strftime('%Y-%m-%d')}",
            content=report_content,
            html_content=markdown_to_html(report_content)
        )
        db.session.add(report)

        # 更新对话标题
        conversation.title = '已完成'
        db.session.commit()

        return jsonify({
            'report_id': report.id,
            'content': report_content
        })

    except Exception as e:
        return jsonify({'error': str(e)}), 500


def markdown_to_html(markdown_text):
    """简单的Markdown转HTML（实际项目中建议用markdown库）"""
    import re

    # 标题
    text = re.sub(r'^### (.+)$', r'<h3>\1</h3>', markdown_text, flags=re.MULTILINE)
    text = re.sub(r'^## (.+)$', r'<h2>\1</h2>', text, flags=re.MULTILINE)
    text = re.sub(r'^# (.+)$', r'<h1>\1</h1>', text, flags=re.MULTILINE)

    # 粗体
    text = re.sub(r'\*\*(.+?)\*\*', r'<strong>\1</strong>', text)

    # 列表
    text = re.sub(r'^- (.+)$', r'<li>\1</li>', text, flags=re.MULTILINE)

    # 段落
    text = re.sub(r'\n\n', '</p><p>', text)
    text = '<p>' + text + '</p>'

    return text


@app.route('/api/report/download/<int:report_id>')
@login_required
def download_report(report_id):
    """下载报告"""
    report = Report.query.get_or_404(report_id)
    if report.user_id != current_user.id:
        return jsonify({'error': '无权访问'}), 403

    # 生成Word文档
    from docx import Document
    from io import BytesIO

    doc = Document()
    doc.add_heading(report.title, 0)

    # 简单解析内容
    lines = report.content.split('\n')
    for line in lines:
        if line.strip():
            if line.startswith('#'):
                doc.add_heading(line.replace('#', '').strip(), level=2)
            elif line.startswith('-'):
                doc.add_paragraph(line.strip(), style='List Bullet')
            else:
                doc.add_paragraph(line.strip())

    # 保存到内存
    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    from flask import send_file
    return send_file(
        buffer,
        as_attachment=True,
        download_name=f"{report.title}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


# ==================== 设置 ====================

@app.route('/settings', methods=['GET', 'POST'])
@login_required
def settings():
    """模型配置页面"""
    if request.method == 'POST':
        current_user.model_provider = request.form.get('provider', 'openai')
        current_user.model_name = request.form.get('model_name', 'gpt-4')
        current_user.api_key = request.form.get('api_key', '')

        db.session.commit()
        flash('配置已保存', 'success')
        return redirect(url_for('settings'))

    return render_template('settings.html',
                         user=current_user,
                         providers=[
                             {'id': 'openai', 'name': 'OpenAI (GPT)'},
                             {'id': 'claude', 'name': 'Claude'},
                             {'id': 'baidu', 'name': '百度文心一言'},
                             {'id': 'alibaba', 'name': '阿里通义千问'},
                             {'id': 'tencent', 'name': '腾讯混元'},
                             {'id': 'zhipu', 'name': '智谱AI (GLM)'},
                             {'id': 'moonshot', 'name': '月之暗面 (Moonshot)'},
                             {'id': 'yi', 'name': '零一万物 (Yi)'},
                             {'id': 'minimax', 'name': 'MiniMax'}
                         ])


# ==================== 管理员API设置 ====================

@app.route('/admin/update-api', methods=['POST'])
def admin_update_api():
    """管理员更新API配置"""
    api_key = request.form.get('api_key', '')
    provider = request.form.get('provider', 'openai')
    api_url = request.form.get('api_url', '')
    model_name = request.form.get('model_name', '')

    # 更新环境变量（仅当前进程生效）
    os.environ['OPENAI_API_KEY'] = api_key
    os.environ['AI_PROVIDER'] = provider
    os.environ['AI_API_URL'] = api_url
    os.environ['AI_MODEL'] = model_name

    # 写入配置文件（永久保存）- 放在项目外部避免更新时被覆盖
    config_path = os.path.join(os.path.dirname(__file__), '..', 'api_config.json')
    config_data = {
        'api_key': api_key,
        'provider': provider,
        'api_url': api_url,
        'model_name': model_name
    }
    with open(config_path, 'w', encoding='utf-8') as f:
        json.dump(config_data, f, ensure_ascii=False, indent=2)

    # 返回到聊天页面，不显示提示
    return redirect(url_for('chat'))


# ==================== 错误处理 ====================

@app.errorhandler(404)
def not_found(e):
    return render_template('error.html', error='页面不存在'), 404


@app.errorhandler(500)
def server_error(e):
    return render_template('error.html', error='服务器错误'), 500


if __name__ == '__main__':
    app.run(debug=True, host='0.0.0.0', port=5000)
